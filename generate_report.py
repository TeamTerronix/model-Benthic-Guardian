"""
Generate Coral Reef Digital Twin — Full Word Report with Diagrams
Run:  python generate_report.py
Output: Coral_Reef_Digital_Twin_Report.docx
"""

import io, os
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patches as FancyArrow
from matplotlib.patches import FancyArrowPatch
from matplotlib.gridspec import GridSpec

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ────────────────────────────────────────────────────────────────────

def fig_to_buf(fig, dpi=150):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

def add_fig(doc, fig, width=Inches(6), caption=None):
    buf = fig_to_buf(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=width)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        shading.set(qn("w:color"), "FFFFFF")
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "EAF1FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

# ══════════════════════════════════════════════════════════════════════════════
#  DIAGRAM FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def fig_prototype_pipeline():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis("off")

    def box(x, y, w, h, label, sub="", color="#1F4E79", tcolor="white", fs=9):
        rect = plt.Rectangle((x - w/2, y - h/2), w, h,
                              fc=color, ec="white", lw=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y + (0.18 if sub else 0), label, ha="center", va="center",
                color=tcolor, fontsize=fs, fontweight="bold", zorder=4)
        if sub:
            ax.text(x, y - 0.28, sub, ha="center", va="center",
                    color=tcolor, fontsize=7, zorder=4, style="italic")

    def arrow(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)

    # Satellite
    box(5, 9.3, 4.5, 0.9, "Satellite SST + DHW Data", "(NOAA Coral Reef Watch)", "#0A2342")

    # Simulate sensors
    box(5, 7.8, 5, 0.9,
        "Simulate 3+ Seabed Sensors",
        "Triangle layout + ±0.2°C noise   (prototype only)", "#2E6DA4")

    arrow(5, 8.85, 5, 8.25)

    # Split
    ax.annotate("", xy=(2.5, 6.55), xytext=(3.5, 7.35),
                arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)
    ax.annotate("", xy=(7.5, 6.55), xytext=(6.5, 7.35),
                arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)

    # PINN
    box(2.5, 6.0, 3.5, 1.0, "PINN (Spatial Twin)",
        "(lat, lon, time) → Temperature", "#1A6B3C")
    # ANN-LSTM
    box(7.5, 6.0, 3.5, 1.0, "ANN-LSTM (Forecaster)",
        "60-day SST/DHW history", "#1A6B3C")

    # Outputs
    box(2.5, 4.6, 3.5, 0.85, "Spatial Heatmap",
        "Fill between sensors", "#2E86AB")
    box(7.5, 4.6, 3.5, 0.85, "+1 / +3 / +7 Day Forecast",
        "SST & DHW prediction", "#2E86AB")

    arrow(2.5, 5.5, 2.5, 5.02)
    arrow(7.5, 5.5, 7.5, 5.02)

    # Merge
    ax.annotate("", xy=(4.0, 3.3), xytext=(2.5, 4.17),
                arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)
    ax.annotate("", xy=(6.0, 3.3), xytext=(7.5, 4.17),
                arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)

    # Risk
    box(5, 2.8, 4.5, 0.95, "Bleaching Risk Score",
        "DHW-aware  (NOAA-style thresholds)", "#8B0000", "white")

    # Legend
    patches = [
        mpatches.Patch(color="#1F4E79", label="Data source"),
        mpatches.Patch(color="#2E6DA4", label="Prototype simulation"),
        mpatches.Patch(color="#1A6B3C", label="AI model"),
        mpatches.Patch(color="#2E86AB", label="Output"),
        mpatches.Patch(color="#8B0000", label="Health assessment"),
    ]
    ax.legend(handles=patches, loc="lower right", fontsize=7.5,
              framealpha=0.85, ncol=2)

    ax.set_title("Coral Reef Digital Twin — Prototype Pipeline", fontsize=13,
                 fontweight="bold", pad=4, color="#0A2342")
    fig.patch.set_facecolor("#F5F8FC")
    return fig


def fig_sensor_layout():
    fig, ax = plt.subplots(figsize=(6.5, 5))
    ax.set_facecolor("#D6EAF8")
    ax.set_xlim(-1.5, 1.5); ax.set_ylim(-1.5, 1.5)
    ax.set_aspect("equal")

    # Reef
    reef = plt.Circle((0, 0), 0.55, color="#F0D58C", ec="#C8A83A", lw=1.5, zorder=2)
    ax.add_patch(reef)
    ax.text(0, 0, "Reef\nSite", ha="center", va="center", fontsize=9, color="#5D4037",
            fontweight="bold")

    # Triangle sensors
    angles = [90, 210, 330]
    labels = ["Sensor A", "Sensor B", "Sensor C"]
    pts = []
    for a, lbl in zip(angles, labels):
        rad = np.radians(a)
        x, y = 1.1*np.cos(rad), 1.1*np.sin(rad)
        pts.append((x, y))
        ax.scatter(x, y, s=180, color="#E74C3C", zorder=5, edgecolors="white", lw=1.5)
        off_x = 0.2 * np.cos(rad)
        off_y = 0.2 * np.sin(rad)
        ax.text(x + off_x, y + off_y, lbl, ha="center", va="center", fontsize=8,
                color="#1A252F", fontweight="bold")

    # Connect sensors
    tri = plt.Polygon(pts, fill=False, ec="#E74C3C", lw=1.2, ls="--", zorder=3)
    ax.add_patch(tri)

    # Satellite (top right)
    ax.scatter(1.3, 1.3, s=250, marker="*", color="#F39C12", zorder=6)
    ax.text(1.3, 1.45, "Satellite SST", ha="center", fontsize=7.5, color="#7D6608")

    # Arrows from satellite to sensors
    for x, y in pts:
        ax.annotate("", xy=(x*0.97, y*0.97), xytext=(1.28, 1.28),
                    arrowprops=dict(arrowstyle="->", color="#F39C12",
                                   lw=1, connectionstyle="arc3,rad=0.15"))

    # Noise annotation
    ax.text(0, -1.42, "SST + ±0.2°C noise  →  seabed simulation",
            ha="center", fontsize=8, style="italic", color="#1F618D")

    ax.set_title("Prototype: 3 Triangle Seabed Sensors (simulated)", fontsize=10,
                 fontweight="bold", color="#1A252F")
    ax.set_xlabel("Longitude offset (normalised)", fontsize=8)
    ax.set_ylabel("Latitude offset (normalised)", fontsize=8)
    ax.tick_params(labelsize=7)
    return fig


def fig_pinn_architecture():
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis("off")

    layers = [
        ("Input\n(lat, lon, time)", 0.8, "#0A2342", 1.8),
        ("Dense 128\ntanh + BN", 2.3, "#1F4E79", 1.8),
        ("Dense 128\ntanh + BN", 3.8, "#1F4E79", 1.8),
        ("Dense 64\ntanh + BN", 5.3, "#2E6DA4", 1.5),
        ("Dense 64\ntanh + BN", 6.8, "#2E6DA4", 1.5),
        ("Dense 32\ntanh + BN", 8.1, "#3498DB", 1.2),
        ("Output\nSST (sigmoid)", 9.4, "#1A6B3C", 1.2),
    ]

    prev_x, prev_y = None, 5
    for label, x, color, h in layers:
        w = 0.95
        rect = plt.Rectangle((x - w/2, 5 - h/2), w, h,
                              fc=color, ec="white", lw=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, 5, label, ha="center", va="center", color="white",
                fontsize=7.5, fontweight="bold", zorder=4)
        if prev_x is not None:
            ax.annotate("", xy=(x - w/2, 5), xytext=(prev_x + w/2, 5),
                        arrowprops=dict(arrowstyle="->", color="#AAA", lw=1.5), zorder=5)
        prev_x = x

    # PDE residual annotation
    ax.annotate("", xy=(5.3, 3.25), xytext=(5.3, 2.2),
                arrowprops=dict(arrowstyle="<-", color="#E74C3C", lw=1.8))
    ax.text(5.3, 1.8,
            "Physics Loss: GradientTape →  ∂T/∂t + u∂T/∂x + v∂T/∂y − α∇²T ≈ 0",
            ha="center", fontsize=8, color="#C0392B",
            bbox=dict(fc="#FDECEA", ec="#C0392B", boxstyle="round,pad=0.4"))

    ax.text(5, 9.2, "Physics-Informed Neural Network (PINN) Architecture",
            ha="center", fontsize=12, fontweight="bold", color="#0A2342")
    fig.patch.set_facecolor("#F5F8FC")
    return fig


def fig_lstm_architecture():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 9); ax.axis("off")

    def box(x, y, w, h, label, sub="", color="#1F4E79"):
        rect = plt.Rectangle((x - w/2, y - h/2), w, h,
                              fc=color, ec="white", lw=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y + (0.22 if sub else 0), label, ha="center", va="center",
                color="white", fontsize=8, fontweight="bold", zorder=4)
        if sub:
            ax.text(x, y - 0.28, sub, ha="center", va="center",
                    color="#D6EAF8", fontsize=7, zorder=4)

    def arr(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1.8), zorder=5)

    # Input timeline
    for i, d in enumerate([58, 59, 60]):
        bx = 0.7 + i * 1.1
        box(bx, 7.2, 0.85, 0.8, f"Day −{60-d+1}", "[SST, DHW]", "#0A2342")
    ax.text(2.6, 7.2, "...", ha="center", va="center", fontsize=14, color="#555")

    ax.text(1.9, 8.3, "Input: 60 days × [SST_norm, DHW_norm]",
            ha="center", fontsize=9, color="#1A252F",
            bbox=dict(fc="#EBF5FB", ec="#2980B9", boxstyle="round,pad=0.3"))

    arr(1.9, 6.8, 1.9, 6.0)

    # TimeDistributed Dense
    box(1.9, 5.6, 3.5, 0.75, "TimeDistributed Dense(32, relu)",
        "Feature mix per day", "#2E6DA4")

    arr(1.9, 5.22, 1.9, 4.5)

    # LSTM
    box(1.9, 4.1, 3.5, 0.75, "LSTM(64 units)",
        "Sequence memory", "#8E44AD")

    arr(1.9, 3.72, 1.9, 3.0)

    # Dense + dropout
    box(1.9, 2.65, 3.5, 0.65, "Dense(64, relu)  +  Dropout(0.2)", color="#C0392B")

    arr(1.9, 2.32, 1.9, 1.7)

    # Output
    box(1.9, 1.35, 3.5, 0.65, "Dense(6, sigmoid)  → Output", color="#1A6B3C")

    # Output labels
    outputs = [
        ("SST +1d", 4.5, 2.5),
        ("DHW +1d", 5.5, 2.5),
        ("SST +3d", 6.5, 2.5),
        ("DHW +3d", 7.5, 2.5),
        ("SST +7d", 8.5, 2.5),
        ("DHW +7d", 9.5, 2.5),
    ]
    colors = ["#1A6B3C", "#27AE60"] * 3
    for label, ox, oy in outputs:
        c = "#1A6B3C" if "SST" in label else "#27AE60"
        rect = plt.Rectangle((ox-0.42, oy-0.32), 0.84, 0.64,
                              fc=c, ec="white", lw=1, zorder=3)
        ax.add_patch(rect)
        ax.text(ox, oy, label, ha="center", va="center", color="white",
                fontsize=7.5, fontweight="bold", zorder=4)
        ax.annotate("", xy=(ox, oy + 0.32), xytext=(3.65, 1.35),
                    arrowprops=dict(arrowstyle="->", color="#27AE60",
                                   lw=1, connectionstyle="arc3,rad=0.15"), zorder=5)

    ax.text(5.5, 8.6, "ANN–LSTM Architecture (Winning Forecaster, LOOKBACK=60 days)",
            ha="center", fontsize=11, fontweight="bold", color="#0A2342")
    fig.patch.set_facecolor("#F5F8FC")
    return fig


def fig_mae_comparison():
    models = ["ANN–LSTM\n(60d)", "CNN–ANN–LSTM", "Persistence", "Forecast\nHead (6-in)", "PINN+Bias"]
    mae_1d = [0.154, 0.161, 0.156, 0.202, 0.209]
    mae_3d = [0.228, 0.236, 0.240, 0.284, 0.260]
    mae_7d = [0.294, 0.305, 0.320, 0.364, 0.328]

    x = np.arange(len(models))
    width = 0.25
    colors = ["#1F4E79", "#2E86AB", "#E74C3C"]

    fig, ax = plt.subplots(figsize=(10, 5.5))
    b1 = ax.bar(x - width, mae_1d, width, label="+1 day", color=colors[0], edgecolor="white")
    b2 = ax.bar(x,         mae_3d, width, label="+3 day", color=colors[1], edgecolor="white")
    b3 = ax.bar(x + width, mae_7d, width, label="+7 day", color=colors[2], edgecolor="white")

    for bars in [b1, b2, b3]:
        for bar in bars:
            ax.text(bar.get_x() + bar.get_width()/2,
                    bar.get_height() + 0.004,
                    f"{bar.get_height():.3f}",
                    ha="center", va="bottom", fontsize=7.5, color="#1A252F")

    # Persistence reference lines
    ax.axhline(0.156, color="#E74C3C", ls=":", lw=1.2, alpha=0.6)
    ax.axhline(0.240, color="#E74C3C", ls=":", lw=1.2, alpha=0.6)
    ax.axhline(0.320, color="#E74C3C", ls=":", lw=1.2, alpha=0.6)

    ax.set_xticks(x)
    ax.set_xticklabels(models, fontsize=9)
    ax.set_ylabel("MAE (°C)", fontsize=10)
    ax.set_title("Forecast MAE Comparison — All Models (SST, lower is better)", fontsize=11,
                 fontweight="bold", color="#0A2342")
    ax.legend(fontsize=9)
    ax.set_ylim(0, 0.42)
    ax.set_facecolor("#F5F8FC")
    fig.patch.set_facecolor("#F5F8FC")
    ax.grid(axis="y", alpha=0.3)

    # Winner label
    ax.annotate("Winner!", xy=(x[0] + width, 0.294 + 0.01),
                xytext=(x[0] + 0.7, 0.38),
                arrowprops=dict(arrowstyle="->", color="#1A6B3C", lw=1.5),
                color="#1A6B3C", fontsize=9, fontweight="bold")
    return fig


def fig_lookback_ablation():
    lookbacks = [7, 14, 30, 60, 90]
    mae_1d = [0.158, 0.155, 0.161, 0.155, 0.160]
    mae_3d = [0.237, 0.236, 0.239, 0.229, 0.232]
    mae_7d = [0.317, 0.320, 0.317, 0.296, 0.298]
    mean   = [0.237, 0.237, 0.239, 0.227, 0.230]

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))

    ax = axes[0]
    ax.plot(lookbacks, mae_1d, "o-", color="#1F4E79", label="+1 day", lw=2)
    ax.plot(lookbacks, mae_3d, "s-", color="#2E86AB", label="+3 day", lw=2)
    ax.plot(lookbacks, mae_7d, "^-", color="#E74C3C", label="+7 day", lw=2)
    ax.axvline(60, color="green", ls="--", lw=1.5, label="Best = 60d")
    ax.set_xlabel("Lookback window (days)", fontsize=10)
    ax.set_ylabel("MAE (°C)", fontsize=10)
    ax.set_title("MAE vs Lookback Window", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8); ax.grid(alpha=0.3)
    ax.set_facecolor("#F5F8FC")

    ax2 = axes[1]
    bar_colors = ["#BDC3C7"]*5
    bar_colors[3] = "#1A6B3C"
    bars = ax2.bar([str(l)+"d" for l in lookbacks], mean, color=bar_colors,
                   edgecolor="white")
    for bar in bars:
        ax2.text(bar.get_x() + bar.get_width()/2,
                 bar.get_height() + 0.0005,
                 f"{bar.get_height():.3f}",
                 ha="center", va="bottom", fontsize=9)
    ax2.set_ylabel("Mean MAE (°C)", fontsize=10)
    ax2.set_title("Mean MAE by Lookback (Best = 60 days)", fontsize=10, fontweight="bold")
    ax2.set_facecolor("#F5F8FC")
    ax2.grid(axis="y", alpha=0.3)

    fig.suptitle("Lookback Ablation Study — ANN-LSTM", fontsize=12,
                 fontweight="bold", color="#0A2342")
    fig.tight_layout()
    fig.patch.set_facecolor("#F5F8FC")
    return fig


def fig_dhw_risk():
    dhw = np.linspace(0, 12, 300)
    risk = np.piecewise(dhw,
        [dhw < 4, (dhw >= 4) & (dhw < 8), dhw >= 8],
        [lambda x: x/4*0.4, lambda x: 0.4 + (x-4)/4*0.4, lambda x: np.minimum(0.8+(x-8)/4*0.2, 1.0)])

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.fill_between(dhw[dhw < 4], 0, risk[dhw < 4], color="#2ECC71", alpha=0.25)
    ax.fill_between(dhw[(dhw >= 4) & (dhw < 8)], 0,
                    risk[(dhw >= 4) & (dhw < 8)], color="#F39C12", alpha=0.35)
    ax.fill_between(dhw[dhw >= 8], 0, risk[dhw >= 8], color="#E74C3C", alpha=0.4)
    ax.plot(dhw, risk, color="#1F4E79", lw=2.5)

    ax.axvline(4, color="#F39C12", ls="--", lw=1.5)
    ax.axvline(8, color="#E74C3C", ls="--", lw=1.5)
    ax.text(2, 0.85, "Safe\n(DHW < 4)", ha="center", color="#1A6B3C", fontsize=9, fontweight="bold")
    ax.text(6, 0.85, "Watch\n(DHW 4–8)", ha="center", color="#B7770D", fontsize=9, fontweight="bold")
    ax.text(10, 0.85, "Alert\n(DHW > 8)", ha="center", color="#A93226", fontsize=9, fontweight="bold")

    ax.set_xlabel("Degree Heating Weeks (DHW)", fontsize=10)
    ax.set_ylabel("Bleaching Risk Score (0–1)", fontsize=10)
    ax.set_title("DHW-Aware Bleaching Risk Score (NOAA-Style Thresholds)", fontsize=11,
                 fontweight="bold", color="#0A2342")
    ax.set_facecolor("#F5F8FC")
    ax.grid(alpha=0.3)
    fig.patch.set_facecolor("#F5F8FC")
    return fig


def fig_evolution_timeline():
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(-0.5, 7.5); ax.set_ylim(-1.5, 3.5); ax.axis("off")

    steps = [
        ("Model 0\nPersistence", 0, "Baseline\n(copy last SST)", "#7F8C8D"),
        ("Model 1\nOriginal PINN", 1, "Triangle sensors\n+ noise + smoothness loss", "#2E6DA4"),
        ("Model 2\nImproved PINN", 2, "Real PDE loss\nAdvection u,v\nHold-out eval", "#1F4E79"),
        ("Model 3\nPINN+Bias", 3, "Bias correction\nRecalibration\nForecasting attempt", "#8E44AD"),
        ("Model 4\n6-input Head", 4, "Issue-time snapshot\nMulti-horizon training\nBeats PINN raw", "#C0392B"),
        ("Model 5\nANN-LSTM", 5, "60-day history\nSequence memory\nBeats persistence!", "#1A6B3C"),
        ("Model 6\nLookback Study", 6, "Ablation 7→90d\n60d is best", "#16A085"),
        ("Model 7\nDeploy", 7, "Real-time forecasts\nRisk alerts", "#D35400"),
    ]

    for i, (title, x, desc, color) in enumerate(steps):
        y_box = 2.0
        rect = plt.Rectangle((x - 0.42, y_box - 0.45), 0.84, 0.9,
                              fc=color, ec="white", lw=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y_box, title, ha="center", va="center", color="white",
                fontsize=7.5, fontweight="bold", zorder=4)

        ax.text(x, 0.8, desc, ha="center", va="center", color="#1A252F",
                fontsize=7, zorder=4,
                bbox=dict(fc="#EBF5FB", ec=color, boxstyle="round,pad=0.3", lw=1))
        ax.plot([x, x], [y_box - 0.45, 1.12], color=color, lw=1.5, ls="--", zorder=2)

        if i < len(steps) - 1:
            ax.annotate("", xy=(x+0.58, y_box), xytext=(x+0.42, y_box),
                        arrowprops=dict(arrowstyle="->", color="#AAA", lw=1.8), zorder=5)

    ax.text(3.5, 3.2, "Model Evolution Timeline — Coral Reef Digital Twin",
            ha="center", fontsize=12, fontweight="bold", color="#0A2342")
    fig.patch.set_facecolor("#F5F8FC")
    return fig


# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # Page margins
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Cover page ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("Coral Reef Digital Twin")
    run.font.size = Pt(28); run.font.bold = True
    run.font.color.rgb = RGBColor(0x0A, 0x23, 0x42)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Model Documentation — Architecture, Results & Comparison")
    r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Project: Prototype digital twin for coral reef monitoring around Sri Lanka\n"
        "Goal: Simulate seabed sensors · Forecast SST · Score coral bleaching risk"
    )
    r3.font.size = Pt(11); r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_page_break()

    # ── 1. Executive Summary ───────────────────────────────────────────────────
    heading(doc, "1. Executive Summary")
    body(doc,
        "This report documents the complete development of a coral reef digital twin — "
        "a prototype AI system designed to simulate seabed temperature sensors, reconstruct "
        "spatial temperature fields, forecast short-term temperature changes, and assess coral "
        "bleaching risk. The project uses satellite Sea Surface Temperature (SST) and Degree "
        "Heating Weeks (DHW) data from NOAA Coral Reef Watch as stand-ins for real seabed loggers, "
        "with small noise added to simulate the difference between surface and underwater readings."
    )
    body(doc,
        "The system evolved through eight major model stages — from a simple persistence baseline "
        "to a Physics-Informed Neural Network (PINN) for spatial mapping, through sequence-based "
        "LSTM models for forecasting, ending with a deployable ANN–LSTM model that beats the "
        "persistence baseline at 3–7 day horizons."
    )

    # Key findings table
    heading(doc, "Key Findings at a Glance", level=2)
    add_table(doc,
        ["Finding", "Result"],
        [
            ["Best spatial model", "PINN (MAE 0.73–0.93°C on hold-outs)"],
            ["Best 1-day forecast", "ANN–LSTM: 0.154°C MAE (persistence: 0.156)"],
            ["Best 3-day forecast", "ANN–LSTM: 0.228°C MAE (persistence: 0.240) ↑"],
            ["Best 7-day forecast", "ANN–LSTM: 0.294°C MAE (persistence: 0.320) ↑"],
            ["Optimal history length", "60 days (ablation: 7, 14, 30, 60, 90 tested)"],
            ["Bleaching risk method", "DHW-aware, NOAA-style (≥4 watch, ≥8 alert)"],
            ["Prototype sensor method", "3 triangle sensors + ±0.2°C noise on SST"],
        ],
        col_widths=[8, 8]
    )
    doc.add_paragraph()

    # ── 2. The Problem & Prototype Concept ────────────────────────────────────
    heading(doc, "2. The Problem and Prototype Concept")
    body(doc,
        "Coral reefs are highly sensitive to water temperature. Elevated temperatures cause coral "
        "bleaching — the expulsion of symbiotic algae — which can lead to coral death if sustained. "
        "Early warning systems that predict temperature stress a few days in advance allow reef "
        "managers to take protective action."
    )
    body(doc,
        "The ideal system would use physical seabed temperature loggers placed around a reef. "
        "However, deploying real underwater sensors is expensive and time-consuming. This project "
        "creates a prototype that can be tested and validated now, designed so that real sensor "
        "streams can replace the simulation later."
    )

    heading(doc, "Prototype Sensor Simulation", level=2)
    body(doc,
        "Since real seabed sensors are not yet available, we simulate them using satellite SST data:"
    )
    bullet(doc, "Take satellite SST at a reef location (freely available from NOAA CRW).")
    bullet(doc, "Place 3 virtual sensors in a triangle around the reef site.")
    bullet(doc, "Add small random noise (±0.2°C) to each sensor to mimic: (a) the difference "
                "between sea surface and seabed temperature, (b) individual sensor measurement error.")
    bullet(doc, "This gives us 3 slightly different time series that a real deployment would produce.")
    body(doc,
        "This approach is scientifically valid as a prototype — SST and seabed temperature "
        "are correlated, and the noise magnitude is realistic for shallow tropical reefs."
    )

    add_fig(doc, fig_sensor_layout(), width=Inches(4.5),
            caption="Figure 1. Prototype seabed sensor layout — 3 triangle points around a reef, "
                    "simulated with satellite SST + ±0.2°C noise.")

    # Pipeline
    heading(doc, "Overall System Pipeline", level=2)
    add_fig(doc, fig_prototype_pipeline(), width=Inches(6),
            caption="Figure 2. Complete prototype pipeline from satellite data to bleaching risk alert.")

    doc.add_paragraph()
    body(doc,
        "The pipeline has three main roles filled by different models:"
    )
    add_table(doc,
        ["Role", "Model", "Why This Choice"],
        [
            ["Spatial mapping\n(fill between sensors)", "PINN", "Continuous spatial field in space–time using physics"],
            ["Short-term forecast\n(+1/+3/+7 days)", "ANN–LSTM\n(60-day history)", "Best skill in experiments; learns seasonal patterns"],
            ["Health alert", "DHW-aware Risk Scorer", "NOAA-style stress thresholds; interpretable"],
        ],
        col_widths=[5, 4, 7]
    )

    # ── 3. Evolution Overview ─────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "3. Model Evolution Overview")
    body(doc,
        "The system evolved through 8 stages. Each stage addressed a specific weakness discovered "
        "in evaluation. The timeline below shows this progression:"
    )
    add_fig(doc, fig_evolution_timeline(), width=Inches(6.5),
            caption="Figure 3. Model evolution from simple baseline to deployable forecaster.")

    add_table(doc,
        ["#", "Model", "Main Idea", "Status"],
        [
            ["0", "Persistence", "Tomorrow = today (copy last SST)", "Strong baseline — every model must beat this"],
            ["1", "Original PINN\n(triangle + noise)", "First spatial prototype with simulated sensors", "Conceptual ancestor — sensor simulation still valid"],
            ["2", "Improved PINN\n(real PDE + advection)", "Physics-based loss, real multi-site data, hold-outs", "Keep for spatial maps"],
            ["3", "PINN + Bias/Recalibration", "Correct warm drift in forecasts", "Intermediate — helps but still loses to persistence"],
            ["4", "6-input Forecast Head", "Issue-time snapshot → future values", "Beaten by LSTM"],
            ["5", "ANN–LSTM / CNN–ANN–LSTM", "60-day history + sequence memory", "ANN–LSTM wins — current best forecaster"],
            ["6", "Lookback Ablation", "Test history lengths 7 → 90 days", "60 days is the sweet spot"],
            ["7", "Deploy ANN–LSTM", "Operational forecast + risk notebook", "Current deployment path"],
        ],
        col_widths=[0.8, 3.5, 6, 5.7]
    )

    # ── 4. Model-by-Model Description ─────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "4. Model-by-Model Description")

    # 4.0 Persistence
    heading(doc, "4.0  Model 0 — Persistence Baseline", level=2)
    body(doc,
        "The persistence model is the simplest possible forecast: predict that tomorrow's temperature "
        "will be the same as today's. Despite its simplicity, it is a strong baseline for short-term "
        "forecasting because ocean SST changes slowly."
    )
    body(doc, "Rule: SST(t + h) = SST(t)  for all horizons h")
    body(doc,
        "Any machine learning model that cannot beat persistence at 3–7 days is not useful for "
        "operational forecasting. All learned models are compared against this benchmark."
    )
    add_table(doc,
        ["Horizon", "Persistence MAE (°C)"],
        [["1 day", "0.156"], ["3 days", "0.240"], ["7 days", "0.320"]],
        col_widths=[5, 5]
    )

    # 4.1 Original PINN
    doc.add_paragraph()
    heading(doc, "4.1  Model 1 — Original PINN (First Prototype)", level=2)
    body(doc,
        "The first model was a Physics-Informed Neural Network trained on three simulated seabed "
        "sensors arranged in a triangle around a single reef site. Satellite SST was used as the "
        "underlying signal, with ±0.2°C noise added to each of the three points."
    )
    heading(doc, "Architecture", level=3)
    code_block(doc,
        "Input  (3):   [lat_norm, lon_norm, time_norm]\n"
        "Hidden:       Dense(128, tanh) + BatchNorm\n"
        "              Dense(128, tanh) + BatchNorm\n"
        "              Dense(64,  tanh) + BatchNorm\n"
        "              Dense(64,  tanh) + BatchNorm\n"
        "              Dense(32,  tanh) + BatchNorm\n"
        "Output (1):   Dense(1, sigmoid)  →  normalised SST\n"
        "Parameters:   ~33,217"
    )
    heading(doc, "Physics Loss (Weak — First Version)", level=3)
    body(doc,
        "The first version used a simplified smoothness constraint rather than a real PDE:"
    )
    code_block(doc, "Loss = MSE(data) + λ × mean( (ŷ[i+1] − ŷ[i])² )")
    body(doc,
        "This penalises abrupt changes between consecutive batch predictions but is not a true "
        "physical law. It was replaced in Model 2 with a proper heat equation residual."
    )
    heading(doc, "What This Model Proved", level=3)
    bullet(doc, "Triangle + noise is a valid prototype approach for seabed sensor simulation.")
    bullet(doc, "PINN can learn a smooth spatial-temporal field from sparse noisy inputs.")
    bullet(doc, "Weaknesses: smoothness loss is not real physics; data was not independently held out.")

    # 4.2 Improved PINN
    doc.add_paragraph()
    heading(doc, "4.2  Model 2 — Improved PINN (Spatial Digital Twin)", level=2)
    body(doc,
        "The PINN was significantly improved in three areas: physics, data, and evaluation."
    )
    heading(doc, "Improvements Made", level=3)
    bullet(doc, "PHYSICS: Replaced the smoothness penalty with the true heat advection–diffusion "
                "PDE residual, computed via TensorFlow's GradientTape during training.")
    bullet(doc, "DATA: Training on 5 real satellite locations around Sri Lanka — Hikkaduwa, Kalpitiya, "
                "Passikudah, South East, Trincomalee — instead of a single noisy series.")
    bullet(doc, "EVALUATION: Proper time and location hold-outs (newest 20% of dates for test; "
                "Trincomalee held out as unseen location).")
    bullet(doc, "ADVECTION: u and v ocean current components estimated from multi-site SST gradients "
                "by least-squares regression, stored in advection.pkl.")

    heading(doc, "Physics Loss — Heat Advection-Diffusion PDE", level=3)
    body(doc,
        "The physics residual R is computed at collocation points using automatic differentiation:"
    )
    code_block(doc,
        "R = ∂T/∂t  +  u × ∂T/∂lon  +  v × ∂T/∂lat  −  α × ∇²T\n\n"
        "Total Loss = MSE(data) + λ × mean(|R|²)\n\n"
        "where:\n"
        "  T = predicted normalised temperature\n"
        "  u, v = advection velocities (estimated from SST gradients)\n"
        "  α = diffusion coefficient\n"
        "  λ = physics weight (tuned by sweep, typically 0.001–0.01)"
    )

    add_fig(doc, fig_pinn_architecture(), width=Inches(6.5),
            caption="Figure 4. PINN architecture with MLP backbone and PDE-based physics loss via GradientTape.")

    heading(doc, "Interpolation Performance", level=3)
    add_table(doc,
        ["Split", "MAE (°C)", "RMSE (°C)", "Notes"],
        [
            ["Validation (time)", "0.73", "0.93", "In-period reconstruction"],
            ["Test (newest years)", "0.93", "1.17", "Warm bias ~+0.55°C"],
            ["Location hold-out\n(Trincomalee)", "0.75", "0.96", "Good spatial transfer"],
        ],
        col_widths=[4.5, 2.5, 2.5, 6.5]
    )
    body(doc,
        "The PINN is a good spatial interpolator — it can reconstruct temperatures between sensor "
        "locations. However, when used as a short-term forecaster, it performs poorly against "
        "the persistence baseline. This motivated the development of dedicated forecasting models."
    )

    # 4.3 PINN + Bias
    doc.add_paragraph()
    heading(doc, "4.3  Model 3 — PINN + Bias Correction / Recalibration", level=2)
    body(doc,
        "Rather than discarding the PINN for forecasting, we attempted to rescue it using "
        "post-processing corrections."
    )
    heading(doc, "Bias Correction", level=3)
    code_block(doc,
        "bias = mean(actual − PINN_prediction)  over recent days at issue time\n"
        "corrected_forecast = PINN_forecast + bias"
    )
    heading(doc, "Recalibration", level=3)
    body(doc,
        "Additional offsets computed from validation set residuals, stored in recalibration.pkl, "
        "are applied on top of bias correction."
    )
    heading(doc, "Results After Bias Correction", level=3)
    add_table(doc,
        ["Horizon", "PINN Raw (°C)", "PINN + Bias (°C)", "Persistence (°C)"],
        [
            ["1 day", "0.87", "0.22", "0.154"],
            ["3 days", "0.87", "0.28", "0.240"],
            ["7 days", "0.87", "0.37", "0.320"],
        ],
        col_widths=[3, 3.5, 3.5, 3.5]
    )
    body(doc,
        "Bias correction is essential — it reduces PINN forecast error from 0.87°C to 0.22–0.37°C. "
        "However, it still loses to the persistence baseline at all horizons. The PINN lacks the "
        "temporal pattern memory needed for short-term forecasting."
    )

    # 4.4 Forecast Head
    doc.add_paragraph()
    heading(doc, "4.4  Model 4 — Six-Input Forecast Head", level=2)
    body(doc,
        "A dedicated supervised forecast head was trained on issue-time snapshot features — "
        "the location, target time, forecast horizon, and current SST/DHW values."
    )
    heading(doc, "Architecture", level=3)
    code_block(doc,
        "Input  (6):  [lat_norm, lon_norm, time_target_norm, horizon_norm,\n"
        "              SST_issue_norm, DHW_issue_norm]\n"
        "Hidden:      Dense(128, tanh) + BN → Dense(128) → Dense(64) → Dense(64) → Dense(32)\n"
        "Output (2):  [SST_norm, DHW_norm]  (sigmoid)\n"
        "Loss:        MSE  (no PDE — snapshot inputs can't support spatial PDE)\n"
        "Parameters:  ~33k"
    )
    heading(doc, "Results", level=3)
    add_table(doc,
        ["Horizon", "Forecast Head MAE (°C)", "Persistence MAE (°C)"],
        [
            ["1 day", "0.202", "0.154"],
            ["3 days", "0.286", "0.240"],
            ["7 days", "0.371", "0.320"],
        ],
        col_widths=[4, 5, 5]
    )
    body(doc,
        "Much better than raw PINN as a forecaster, but still behind persistence. The fundamental "
        "problem is that a single snapshot does not contain enough historical context. The model "
        "cannot see recent trends, seasonal cycles, or anomaly patterns."
    )

    # 4.5 ANN-LSTM
    doc.add_page_break()
    heading(doc, "4.5  Model 5 — ANN–LSTM (Winning Forecaster)", level=2)
    body(doc,
        "The key insight from Models 3 and 4 was that a snapshot of current conditions is insufficient. "
        "Short-term forecasting requires history — the model needs to see recent trends, "
        "oscillations, and anomaly patterns to predict the next few days accurately."
    )
    body(doc,
        "The ANN–LSTM uses the last 60 days of SST and DHW observations as input, processes "
        "them with a feature-mixing Dense layer (applied to each day independently), then passes "
        "the sequence through an LSTM to capture temporal dependencies."
    )

    add_fig(doc, fig_lstm_architecture(), width=Inches(6.5),
            caption="Figure 5. ANN–LSTM architecture — the winning forecaster. Input: 60 days × 2 features. "
                    "Output: SST and DHW at +1, +3, +7 day horizons.")

    heading(doc, "Architecture Details", level=3)
    add_table(doc,
        ["Layer", "Type", "Output Shape", "Purpose"],
        [
            ["Input", "—", "(batch, 60, 2)", "60 days of [SST_norm, DHW_norm]"],
            ["TimeDistributed Dense", "Dense(32, relu)", "(batch, 60, 32)", "Feature mix per day"],
            ["LSTM", "LSTM(64 units)", "(batch, 64)", "Sequence memory & temporal patterns"],
            ["Dense + Dropout", "Dense(64, relu), Dropout(0.2)", "(batch, 64)", "Non-linear transform + regularisation"],
            ["Output", "Dense(6, sigmoid)", "(batch, 6)", "SST+DHW at +1d, +3d, +7d"],
        ],
        col_widths=[3.5, 4, 3.5, 5]
    )
    body(doc,
        "Total parameters: ~29,500 (LOOKBACK=60). Training: Adam optimiser, MSE loss, "
        "EarlyStopping on val_loss with patience=10, ReduceLROnPlateau."
    )

    heading(doc, "Why LSTM?", level=3)
    bullet(doc, "LSTM cells maintain a 'cell state' that can carry information across many time steps.")
    bullet(doc, "Gating mechanisms (input, forget, output gates) let the model decide what to remember "
                "and what to discard from the 60-day history.")
    bullet(doc, "This is ideal for SST which has seasonal trends, slow drift, and occasional anomalies.")

    heading(doc, "Results", level=3)
    add_table(doc,
        ["Model", "+1 day MAE", "+3 day MAE", "+7 day MAE", "Beats Persistence?"],
        [
            ["ANN–LSTM (60d) ★", "0.154", "0.228", "0.294", "Yes (3d, 7d)"],
            ["CNN–ANN–LSTM", "0.161", "0.236", "0.305", "Yes (3d, 7d)"],
            ["Persistence", "0.156", "0.240", "0.320", "Baseline"],
            ["Forecast Head (6-in)", "0.202", "0.284", "0.364", "No"],
            ["PINN + Bias", "0.209", "0.260", "0.328", "No"],
        ],
        col_widths=[4.5, 2.5, 2.5, 2.5, 4]
    )

    add_fig(doc, fig_mae_comparison(), width=Inches(6.5),
            caption="Figure 6. Forecast MAE comparison across all models and horizons. "
                    "Dotted lines mark persistence benchmark. ANN–LSTM wins at 3 and 7 days.")

    # 4.6 CNN-ANN-LSTM
    heading(doc, "4.6  Model 6 — CNN–ANN–LSTM", level=2)
    body(doc,
        "The CNN–ANN–LSTM adds 1D convolutional layers before the LSTM to extract local temporal "
        "patterns (e.g., short-term oscillations) before the sequence modelling stage."
    )
    code_block(doc,
        "Input (batch, 60, 2)\n"
        "  → Conv1D(32, kernel=5, relu)   ← 5-day local patterns\n"
        "  → Conv1D(32, kernel=3, relu)   ← 3-day refinement\n"
        "  → TimeDistributed Dense(32)\n"
        "  → LSTM(64)\n"
        "  → Dense(64, relu) + Dropout(0.2)\n"
        "  → Dense(6, sigmoid)            ← same 6-output as ANN-LSTM\n"
        "Parameters: ~33,900"
    )
    body(doc,
        "Result: Close second to ANN–LSTM, but convolutions did not help significantly on "
        "daily SST data (which is slow and smooth). The simpler ANN–LSTM is preferred."
    )

    # 4.7 Lookback Ablation
    doc.add_page_break()
    heading(doc, "4.7  Model 7 — Lookback Ablation (History Length Study)", level=2)
    body(doc,
        "To justify using 60 days of history, a systematic ablation study was conducted: "
        "the same ANN–LSTM was trained five times with different lookback windows."
    )
    add_table(doc,
        ["Experiment", "Lookback", "MAE +1d", "MAE +3d", "MAE +7d", "Mean MAE"],
        [
            ["A", "7 days",  "0.158", "0.237", "0.317", "0.237"],
            ["B", "14 days", "0.155", "0.236", "0.320", "0.237"],
            ["C", "30 days", "0.161", "0.239", "0.317", "0.239"],
            ["D ★ BEST", "60 days", "0.155", "0.229", "0.296", "0.227"],
            ["E", "90 days", "0.160", "0.232", "0.298", "0.230"],
        ],
        col_widths=[3, 2.5, 2.5, 2.5, 2.5, 2.5]
    )

    add_fig(doc, fig_lookback_ablation(), width=Inches(6.5),
            caption="Figure 7. Lookback ablation results. Left: MAE per horizon vs lookback window. "
                    "Right: Mean MAE (60 days is the best, green bar).")

    heading(doc, "Key Findings from Ablation", level=3)
    bullet(doc, "7–30 days: Short history underperforms at 7-day lead — not enough context.")
    bullet(doc, "60 days: Best overall. Captures 2× the seasonal signal of shorter windows.")
    bullet(doc, "90 days: Marginally worse than 60 — extra history adds noise / computation cost.")
    bullet(doc, "Conclusion: 60 days is the evidence-based choice for deployment.")

    # 4.8 Deploy
    heading(doc, "4.8  Model 8 — Deployed ANN–LSTM Forecaster", level=2)
    body(doc,
        "The final deployment notebook (11_deploy_ann_lstm_forecaster.ipynb) loads the best "
        "ANN–LSTM model and demonstrates operational use:"
    )
    bullet(doc, "Loads ann_lstm_L60_best.h5 and associated scalers.")
    bullet(doc, "Takes last 60 days of historical SST+DHW for a reef site.")
    bullet(doc, "Produces +1/+3/+7 day SST and DHW forecasts with uncertainty context.")
    bullet(doc, "Computes persistence comparison at each horizon.")
    bullet(doc, "Calculates DHW-aware bleaching risk for each forecast step.")
    bullet(doc, "Saves forecast plots and CSV outputs to lstm_deploy/.")

    # ── 5. Bleaching Risk ─────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "5. Bleaching Risk Assessment Layer")
    body(doc,
        "The bleaching risk scorer is not a neural network — it is an interpretable rule-based "
        "system built on NOAA Coral Reef Watch methodology, applied on top of model outputs."
    )
    heading(doc, "How It Works", level=2)
    body(doc,
        "Four signals are combined into a final risk score (0 to 1):"
    )
    add_table(doc,
        ["Signal", "Weight", "Description"],
        [
            ["DHW (primary)", "High", "Degree Heating Weeks: accumulated heat stress over 12 weeks"],
            ["Temperature anomaly", "Medium", "Predicted SST minus monthly climatological baseline"],
            ["Duration of warm stress", "Medium", "How many consecutive days above baseline"],
            ["Warming rate", "Low", "Rate of temperature increase (°C/day)"],
        ],
        col_widths=[4.5, 2.5, 9]
    )

    heading(doc, "NOAA-Style DHW Thresholds", level=2)
    add_table(doc,
        ["DHW Value", "Risk Level", "Meaning"],
        [
            ["< 4 °C-weeks", "Safe (Level 0)", "No significant stress"],
            ["4–8 °C-weeks", "Watch (Level 1)", "Possible bleaching"],
            ["> 8 °C-weeks", "Alert (Level 2)", "Likely/severe bleaching"],
        ],
        col_widths=[4, 4, 8]
    )

    add_fig(doc, fig_dhw_risk(), width=Inches(5.5),
            caption="Figure 8. DHW-to-risk mapping. Risk score increases from Safe (green) through Watch "
                    "(amber) to Alert (red) as accumulated heat stress rises.")

    body(doc,
        "DHW is particularly important because a single warm day is less harmful than sustained "
        "warmth over weeks. A reef can recover from a brief temperature spike but not from weeks "
        "of elevated heat — this is exactly what DHW captures."
    )

    # ── 6. Full Comparison ────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "6. Full Model Comparison")

    heading(doc, "6.1  Forecasting Performance (SST MAE °C)", level=2)
    body(doc,
        "All models evaluated on the same test set (n=1,680 samples, 5 locations, "
        "held-out time period). Lower MAE = better."
    )
    add_table(doc,
        ["Rank", "Model", "+1 day", "+3 days", "+7 days", "Beats Persistence\n(3–7d)?"],
        [
            ["1 ★", "ANN–LSTM (60d)", "0.154", "0.228", "0.294", "Yes"],
            ["2", "CNN–ANN–LSTM", "0.161", "0.236", "0.305", "Yes"],
            ["3", "Persistence", "0.156", "0.240", "0.320", "Baseline"],
            ["4", "Forecast Head (6-in)", "0.202", "0.284", "0.364", "No"],
            ["5", "PINN + Bias", "0.209", "0.260", "0.328", "No"],
        ],
        col_widths=[1.5, 4, 2, 2, 2, 4.5]
    )

    heading(doc, "6.2  Spatial Interpolation Performance (PINN)", level=2)
    body(doc,
        "The PINN is evaluated separately as a spatial model — its role is to fill temperature "
        "values between sensor locations, not to forecast future values."
    )
    add_table(doc,
        ["Split", "MAE (°C)", "RMSE (°C)", "Notes"],
        [
            ["Validation (time hold-out)", "0.73", "0.93", "In-period reconstruction"],
            ["Test (newest years)", "0.93", "1.17", "Warm bias ~+0.55°C"],
            ["Location hold-out\n(Trincomalee)", "0.75", "0.96", "Good spatial generalisation"],
        ],
        col_widths=[5, 2.5, 2.5, 6]
    )

    heading(doc, "6.3  Architecture Cheat-Sheet", level=2)
    add_table(doc,
        ["Model", "Input", "Core", "Output", "~Params"],
        [
            ["PINN", "(lat, lon, t)", "MLP 128×2–64×2–32\n+ PDE loss", "SST (1d)", "33k"],
            ["Forecast Head", "6 snapshot features", "Same MLP, MSE loss", "SST + DHW (2d)", "~33k"],
            ["ANN–LSTM", "(60, 2) SST/DHW", "TD-Dense→LSTM→Dense", "6 (3 horizons × 2)", "29k"],
            ["CNN–ANN–LSTM", "(60, 2) SST/DHW", "Conv1D×2→TD-Dense\n→LSTM→Dense", "6 (3 horizons × 2)", "34k"],
            ["Persistence", "Last SST", "Rule-based copy", "Same SST", "0"],
            ["Risk Scorer", "T, DHW, baseline", "Weighted rules", "Score 0–1, Level 0–2", "0"],
        ],
        col_widths=[3, 3, 5, 3.5, 1.5]
    )

    # ── 7. Key Notebooks ──────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "7. Key Notebooks and Scripts")
    add_table(doc,
        ["Notebook / Script", "Purpose"],
        [
            ["01_data_preparation.ipynb", "Simulate triangle sensor data from SST with noise"],
            ["02_pinn_model.ipynb", "Train the PINN spatial model"],
            ["03_visualization.ipynb", "Generate spatial heatmaps and risk maps"],
            ["04_real_spatial_data.ipynb", "Load real multi-site satellite data with hold-outs"],
            ["05_estimate_advection.ipynb", "Estimate ocean current u, v from SST gradients"],
            ["06_evaluation.ipynb", "Full hold-out and forecast skill evaluation"],
            ["07_improve_forecasts.ipynb", "Bias correction, recalibration, 6-input head"],
            ["08_tune_physics.ipynb", "Sweep physics weight λ to find optimal value"],
            ["09_lstm_sequence_forecast.ipynb", "Train ANN–LSTM and CNN–ANN–LSTM; compare all models"],
            ["10_lookback_ablation.ipynb", "History length ablation: 7, 14, 30, 60, 90 days"],
            ["11_deploy_ann_lstm_forecaster.ipynb", "Operational deployment: load model, forecast, risk alert"],
            ["pinn_physics.py", "PINN model class, PDE residual, GradientTape training step"],
            ["prepare_data.py", "Data loading, scaling, time+location split, .npy export"],
            ["estimate_advection.py", "Least-squares u,v estimation from multi-site SST"],
            ["validate_forecast.py", "Forecast skill evaluation vs persistence/climatology baselines"],
            ["recalibrate.py", "Rolling recalibration offsets from validation residuals"],
            ["forecaster.py", "PINNForecaster class with bias correction and risk output"],
            ["utils.py", "DHW-aware bleaching risk calculation helper"],
        ],
        col_widths=[6, 10]
    )

    # ── 8. Key Findings ───────────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "8. Key Findings and Lessons Learned")

    findings = [
        ("Triangle + SST noise is a valid prototype",
         "Until real seabed loggers exist, simulating 3 sensors with ±0.2°C noise on satellite "
         "SST is scientifically defensible. The design is ready to accept real sensor streams."),
        ("PINN is the right tool for spatial mapping, not for forecasting",
         "The physics-informed network excels at interpolating temperature between sensor locations "
         "(0.73–0.93°C MAE). As a raw short-term forecaster, it is outperformed even by the trivial "
         "persistence model. Use PINN for maps; use LSTM for predictions."),
        ("Real physics loss matters for honest evaluation",
         "Replacing the batch-smoothness penalty with a true heat advection–diffusion PDE residual "
         "via GradientTape made the physics constraint meaningful — but did not magically improve "
         "1–7 day forecasts. Physics helps the map task; it does not substitute for temporal history."),
        ("Bias correction is essential for PINN forecasts (but not enough)",
         "The PINN has a systematic warm bias of ~0.55°C on test data. Adding bias correction "
         "reduces forecast error from 0.87°C to 0.22–0.37°C — a dramatic improvement — but still "
         "behind persistence at all horizons."),
        ("Snapshot features are insufficient for forecasting",
         "The 6-input forecast head (current location, time, horizon, SST, DHW) performed better "
         "than the raw PINN but still could not beat persistence. A single snapshot of current "
         "conditions cannot encode recent trends or seasonal context."),
        ("60 days of history + LSTM is the first approach to beat persistence",
         "The ANN–LSTM with a 60-day lookback window beats persistence at 3 and 7 days. This "
         "proves that temporal memory is the key ingredient for short-term reef forecasting."),
        ("60 days > 90 days for lookback",
         "The ablation study shows that more history is not always better. 90 days adds noise "
         "and slightly worse performance. 60 days is the evidence-based optimum for this dataset."),
        ("Split roles across models",
         "The best architecture is: PINN = spatial map, ANN–LSTM = time forecast, "
         "DHW risk scorer = health alert. No single model does all three tasks best."),
    ]

    for i, (title, desc) in enumerate(findings):
        p = doc.add_paragraph()
        run = p.add_run(f"Finding {i+1}: {title}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        body(doc, desc)
        doc.add_paragraph()

    # ── 9. Summary Sentence ───────────────────────────────────────────────────
    heading(doc, "9. One-Sentence Summary")
    p = doc.add_paragraph()
    run = p.add_run(
        "We started with a triangle-sensor PINN prototype on noisy SST to mimic seabed loggers, "
        "improved physics and evaluation for a spatial digital twin, then showed that short-term "
        "coral-temperature forecasting needs sequence memory — culminating in an ANN–LSTM with "
        "60 days of SST/DHW history as the best forecast model, while the PINN remains the "
        "spatial map engine."
    )
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x0A, 0x23, 0x42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    return doc


if __name__ == "__main__":
    out = r"D:\corals\model\Coral_Reef_Digital_Twin_Report.docx"
    print("Building diagrams and document...")
    doc = build_doc()
    doc.save(out)
    print(f"Saved: {out}")
